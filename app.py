import streamlit as st
import json
from io import BytesIO
from docx import Document

st.set_page_config(page_title="Mémoire Builder", layout="wide")

st.title("🧠 Mémoire Builder")
st.write("Crée, sauvegarde et recharge la structure de ton mémoire (sujets, paragraphes, exemples, références).")

# --- Initialisation du state ---
if "paragraphs" not in st.session_state:
    st.session_state.paragraphs = []

# --- Upload d'un ancien JSON ---
st.sidebar.header("📂 Charger un projet")
uploaded_file = st.sidebar.file_uploader("Importer un fichier JSON existant", type=["json"])
if uploaded_file is not None:
    data = json.load(uploaded_file)
    if isinstance(data, list):
        st.session_state.paragraphs = data
        st.sidebar.success(f"{len(data)} paragraphes importés ✅")
    else:
        st.sidebar.error("Fichier JSON invalide.")

# --- Fonction pour ajouter un paragraphe ---
def add_paragraph():
    st.session_state.paragraphs.append({
        "sujet": "",
        "texte": "",
        "exemples": "",
        "references": ""
    })

st.sidebar.header("✏️ Édition")
st.sidebar.button("➕ Ajouter un paragraphe", on_click=add_paragraph)

# --- Affichage des paragraphes existants ---
st.write("### Paragraphes actuels")

if len(st.session_state.paragraphs) == 0:
    st.info("Aucun paragraphe pour l'instant. Clique sur ➕ pour commencer.")
else:
    for i, para in enumerate(st.session_state.paragraphs):
        with st.expander(f"Paragraphe {i+1} : {para['sujet'] or 'sans sujet'}", expanded=True):
            st.session_state.paragraphs[i]["sujet"] = st.text_input(
                f"Sujet-clé (#{i+1})", para["sujet"], key=f"sujet_{i}"
            )
            st.session_state.paragraphs[i]["texte"] = st.text_area(
                f"Texte principal (#{i+1})", para["texte"], height=150, key=f"txt_{i}"
            )
            st.session_state.paragraphs[i]["exemples"] = st.text_area(
                f"Exemples concrets (#{i+1})", para["exemples"], height=100, key=f"ex_{i}"
            )
            st.session_state.paragraphs[i]["references"] = st.text_area(
                f"Références APA7 (#{i+1})", para["references"], height=100, key=f"ref_{i}"
            )
            if st.button("🗑️ Supprimer ce paragraphe", key=f"del_{i}"):
                st.session_state.paragraphs.pop(i)
                st.experimental_rerun()

# --- Téléchargement JSON ---
st.sidebar.header("💾 Sauvegarde JSON")
if st.sidebar.button("📥 Générer le JSON"):
    json_data = json.dumps(st.session_state.paragraphs, indent=2, ensure_ascii=False)
    st.sidebar.download_button(
        label="⬇️ Télécharger JSON",
        data=json_data,
        file_name="memoire_data.json",
        mime="application/json"
    )

# --- Export DOCX ---
def generate_docx(paragraphs):
    doc = Document()
    doc.add_heading("Structure du mémoire", level=0)
    for i, p in enumerate(paragraphs, start=1):
        doc.add_heading(f"Paragraphe {i} : {p['sujet']}", level=1)
        if p["texte"]:
            doc.add_paragraph(p["texte"])
        if p["exemples"]:
            doc.add_paragraph(f"🧩 Exemples : {p['exemples']}")
        if p["references"]:
            doc.add_paragraph(f"📚 Références : {p['references']}")
        doc.add_paragraph("")  # ligne vide
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

st.sidebar.header("📝 Export DOCX")
if st.sidebar.button("📤 Générer le DOCX"):
    if st.session_state.paragraphs:
        buffer = generate_docx(st.session_state.paragraphs)
        st.sidebar.download_button(
            label="⬇️ Télécharger le DOCX",
            data=buffer,
            file_name="memoire_structure.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    else:
        st.sidebar.warning("Aucun paragraphe à exporter.")

# --- Aperçu JSON brut ---
with st.expander("🧾 Aperçu du JSON généré"):
    st.json(st.session_state.paragraphs)
